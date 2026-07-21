#!/usr/bin/env python3
"""Build the initial Relay Studio Word library and editable Visio UML atlas.

The generated DOCX files are the authoritative narrative documents. This
bootstrap script exists to make the initial migration reproducible and to keep
the shared formatting and diagram package construction reviewable.
"""

from __future__ import annotations

import json
import math
import re
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence
from xml.sax.saxutils import escape

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = ROOT / "documentation"
WORD_DIR = DOC_ROOT / "word"
VISIO_DIR = DOC_ROOT / "uml" / "visio"
PREVIEW_DIR = DOC_ROOT / "uml" / "previews"
TRACEABILITY_PATH = DOC_ROOT / "documentation-traceability.json"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5F6B7A"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"


@dataclass(frozen=True)
class DiagramNode:
    key: str
    label: str
    category: str = "component"


@dataclass(frozen=True)
class DiagramEdge:
    source: str
    target: str
    label: str = ""


@dataclass(frozen=True)
class DiagramSpec:
    slug: str
    title: str
    uml_type: str
    purpose: str
    nodes: tuple[DiagramNode, ...]
    edges: tuple[DiagramEdge, ...]
    notes: tuple[str, ...]


@dataclass(frozen=True)
class SectionSpec:
    heading: str
    paragraphs: tuple[str, ...] | str = ()
    bullets: tuple[str, ...] = ()
    steps: tuple[str, ...] = ()
    code: tuple[str, ...] = ()
    page_break_before: bool = False


RETAINED_MARKDOWN = {
    "documentation/README.md",
    "documentation/detailed-sprint-plan.md",
    "documentation/build-phase-plan.md",
    "documentation/product-terminology-glossary.md",
    "documentation/secret-redaction-policy.md",
    "documentation/tauri-security-checklist.md",
    "documentation/platform-guidelines-compliance-backlog.md",
    "documentation/live-rest-acceptance-test-matrix.md",
    "documentation/sample-test-project-definition.md",
    "documentation/visual-target.md",
    "documentation/sprint-1-ux-blueprint.md",
    "documentation/current-human-test-script.md",
    "documentation/sprint-10b-3-windows-qa-script.md",
    "documentation/reviews/sprint-18a/review-report.md",
    "documentation/reviews/sprint-18a/remediation-register.md",
    "documentation/reviews/sprint-18b/closure-report.md",
}


DIAGRAMS = (
    DiagramSpec(
        "class-diagram",
        "Relay Studio Project Domain Model",
        "Class Diagram",
        "Defines the persisted project types and their cardinality so schema, import, persistence, and execution changes begin from the same model.",
        (
            DiagramNode("project", "RelayProject\n+services[]\n+flows[]\n+environments[]\n+settings"),
            DiagramNode("service", "ProjectService\nmethod\npath\nauthProfile\nretry"),
            DiagramNode("flow", "ProjectFlow\nnodes[]\nedges[]\nmappings[]"),
            DiagramNode("environment", "ProjectEnvironment\nvariables[]"),
            DiagramNode("response", "SavedResponseMetadata\nstatus\nfilePath\nredacted"),
            DiagramNode("settings", "ProjectSettings\nproxy\ntimeout\ntheme"),
        ),
        (
            DiagramEdge("project", "service", "1 to many"),
            DiagramEdge("project", "flow", "1 to many"),
            DiagramEdge("project", "environment", "1 to many"),
            DiagramEdge("project", "response", "1 to many"),
            DiagramEdge("project", "settings", "owns"),
            DiagramEdge("flow", "service", "references"),
        ),
        ("Source: src/project/projectModel.ts", "Schema version is explicit and validated before use."),
    ),
    DiagramSpec(
        "object-diagram",
        "Sample API Regression Project Snapshot",
        "Object Diagram",
        "Shows representative runtime instances for debugging project load, selection, environment substitution, and saved response behavior.",
        (
            DiagramNode("project", "project:RelayProject\nname=Sample API Regression"),
            DiagramNode("svc1", "createOrder:ProjectService\nmethod=POST"),
            DiagramNode("svc2", "getUser:ProjectService\nmethod=GET"),
            DiagramNode("flow", "orderFlow:ProjectFlow\nstatus=idle"),
            DiagramNode("qa", "qa:ProjectEnvironment\nbaseUrl={{QA_URL}}"),
            DiagramNode("saved", "orderResponse:SavedResponse\nredacted=true"),
        ),
        (
            DiagramEdge("project", "svc1", "contains"),
            DiagramEdge("project", "svc2", "contains"),
            DiagramEdge("project", "flow", "contains"),
            DiagramEdge("project", "qa", "active environment"),
            DiagramEdge("svc1", "saved", "produced"),
        ),
        ("Values are illustrative and contain no credentials.", "Use this snapshot when explaining object identity versus persisted type definitions."),
    ),
    DiagramSpec(
        "component-diagram",
        "Relay Studio Component Architecture",
        "Component Diagram",
        "Shows the major replaceable components and the interfaces across the React, TypeScript, Tauri, filesystem, and network boundaries.",
        (
            DiagramNode("ui", "React Workbench UI", "ui"),
            DiagramNode("designer", "Service and Flow Designers"),
            DiagramNode("runner", "Request and Flow Runners"),
            DiagramNode("project", "Project Persistence"),
            DiagramNode("importer", "OpenAPI Importer"),
            DiagramNode("diagnostics", "Diagnostics and Redaction"),
            DiagramNode("tauri", "Rust Tauri Commands", "native"),
            DiagramNode("fs", "Local Filesystem", "external"),
            DiagramNode("api", "External REST APIs", "external"),
        ),
        (
            DiagramEdge("ui", "designer", "edits"), DiagramEdge("ui", "runner", "executes"),
            DiagramEdge("ui", "project", "open/save"), DiagramEdge("ui", "importer", "imports"),
            DiagramEdge("runner", "diagnostics", "records redacted events"),
            DiagramEdge("project", "tauri", "native IO"), DiagramEdge("runner", "tauri", "native HTTP"),
            DiagramEdge("tauri", "fs", "atomic files"), DiagramEdge("tauri", "api", "HTTP(S)"),
        ),
        ("Browser-mode adapters support automated UI testing.", "Native commands are explicitly allow-listed in the Tauri capability."),
    ),
    DiagramSpec(
        "deployment-diagram",
        "Relay Studio Deployment Topology",
        "Deployment Diagram",
        "Maps unsigned desktop packages and their runtime dependencies across supported operating systems and GitHub Actions.",
        (
            DiagramNode("ci", "GitHub Actions\nquality, security, packaging", "external"),
            DiagramNode("mac", "macOS\nDMG + WebView", "native"),
            DiagramNode("win", "Windows\nMSI/NSIS + WebView2", "native"),
            DiagramNode("linux", "Linux\nDEB/AppImage + WebKitGTK", "native"),
            DiagramNode("files", "User Project and Backup Files", "external"),
            DiagramNode("rest", "Configured REST Targets", "external"),
            DiagramNode("registry", "npm / crates.io / RustSec", "external"),
        ),
        (
            DiagramEdge("ci", "mac", "build artifact"), DiagramEdge("ci", "win", "build artifact"),
            DiagramEdge("ci", "linux", "build artifact"), DiagramEdge("ci", "registry", "dependency metadata"),
            DiagramEdge("mac", "files", "local IO"), DiagramEdge("win", "files", "local IO"),
            DiagramEdge("linux", "files", "local IO"), DiagramEdge("mac", "rest", "HTTP(S)"),
            DiagramEdge("win", "rest", "HTTP(S)"), DiagramEdge("linux", "rest", "HTTP(S)"),
        ),
        ("Packages are intentionally unsigned developer betas.", "Credentials remain local or in protected GitHub secrets."),
    ),
    DiagramSpec(
        "package-diagram",
        "Relay Studio Source Package Dependencies",
        "Package Diagram",
        "Provides a module-level map for locating code before debugging or adding a feature.",
        (
            DiagramNode("app", "src/App.tsx", "ui"),
            DiagramNode("services", "src/services"),
            DiagramNode("project", "src/project"),
            DiagramNode("shell", "src/shell"),
            DiagramNode("lib", "src/lib"),
            DiagramNode("tests", "src/test + e2e"),
            DiagramNode("rust", "src-tauri/src", "native"),
            DiagramNode("docs", "documentation", "external"),
        ),
        (
            DiagramEdge("app", "services", "uses"), DiagramEdge("app", "project", "uses"),
            DiagramEdge("app", "shell", "uses"), DiagramEdge("services", "lib", "redaction/errors"),
            DiagramEdge("project", "lib", "redaction/errors"), DiagramEdge("shell", "rust", "invoke/events"),
            DiagramEdge("tests", "app", "verifies"), DiagramEdge("tests", "services", "verifies"),
            DiagramEdge("docs", "tests", "defines acceptance"),
        ),
        ("Dependencies should point inward toward typed domain and utility modules.", "Avoid hidden coupling through globals or silent fallbacks."),
    ),
    DiagramSpec(
        "profile-diagram",
        "Relay Studio UML Profile",
        "Profile Diagram",
        "Defines project-specific stereotypes used consistently across the architecture atlas.",
        (
            DiagramNode("ui", "<<UI>>\nReact interaction surface", "ui"),
            DiagramNode("service", "<<Service>>\nTyped application logic"),
            DiagramNode("native", "<<NativeCommand>>\nRust/Tauri boundary", "native"),
            DiagramNode("persisted", "<<Persisted>>\nProject or response data"),
            DiagramNode("sensitive", "<<Sensitive>>\nMust be redacted", "sensitive"),
            DiagramNode("external", "<<External>>\nNetwork or filesystem", "external"),
        ),
        (
            DiagramEdge("ui", "service", "invokes"), DiagramEdge("service", "native", "may cross"),
            DiagramEdge("service", "persisted", "creates/reads"), DiagramEdge("sensitive", "persisted", "restricted field"),
            DiagramEdge("native", "external", "accesses"),
        ),
        ("Stereotypes are documentation conventions, not runtime annotations.", "Sensitive data must pass the shared redaction policy."),
    ),
    DiagramSpec(
        "composite-structure-diagram",
        "Request Execution Composite Structure",
        "Composite Structure Diagram",
        "Shows the collaborating parts inside a request execution and their ports to project state, native transport, cancellation, and diagnostics.",
        (
            DiagramNode("context", "Execution Context\nenvironment + variables"),
            DiagramNode("builder", "Request Builder"),
            DiagramNode("retry", "Retry Controller"),
            DiagramNode("abort", "AbortSignal Port"),
            DiagramNode("transport", "Transport Port", "native"),
            DiagramNode("formatter", "Response Formatter"),
            DiagramNode("diagnostics", "Redacted Diagnostics"),
        ),
        (
            DiagramEdge("context", "builder", "resolve variables"), DiagramEdge("builder", "retry", "prepared request"),
            DiagramEdge("abort", "retry", "cancel"), DiagramEdge("retry", "transport", "attempt"),
            DiagramEdge("transport", "formatter", "response"), DiagramEdge("retry", "diagnostics", "attempt events"),
            DiagramEdge("formatter", "diagnostics", "safe summary"),
        ),
        ("Cancellation is cooperative through AbortSignal.", "Retry count and backoff are bounded by the service definition."),
    ),
    DiagramSpec(
        "use-case-diagram",
        "Relay Studio User and Maintainer Use Cases",
        "Use Case Diagram",
        "Explains the value of Relay Studio without requiring source-code knowledge.",
        (
            DiagramNode("developer", "Developer", "external"),
            DiagramNode("qa", "QA Engineer", "external"),
            DiagramNode("maintainer", "Maintainer", "external"),
            DiagramNode("project", "Create/Open Project", "ui"),
            DiagramNode("import", "Import Selected API Operations", "ui"),
            DiagramNode("request", "Build and Run Request", "ui"),
            DiagramNode("flow", "Build and Debug Flow", "ui"),
            DiagramNode("diagnose", "Export Redacted Diagnostics", "ui"),
            DiagramNode("release", "Verify and Package Beta", "ui"),
        ),
        (
            DiagramEdge("developer", "project"), DiagramEdge("developer", "import"),
            DiagramEdge("developer", "request"), DiagramEdge("developer", "flow"),
            DiagramEdge("developer", "diagnose"), DiagramEdge("qa", "request"),
            DiagramEdge("qa", "flow"), DiagramEdge("qa", "diagnose"),
            DiagramEdge("maintainer", "release"), DiagramEdge("maintainer", "diagnose"),
        ),
        ("The product is a personal developer tool, not a hosted collaboration service.", "All imports require explicit operation selection."),
    ),
    DiagramSpec(
        "activity-diagram",
        "Selective OpenAPI Import Activity",
        "Activity Diagram",
        "Models the user-visible decisions and failure paths from a Swagger UI or direct definition URL to committed project services.",
        (
            DiagramNode("start", "Enter definition URL", "ui"),
            DiagramNode("fetch", "Fetch page or definition"),
            DiagramNode("discover", "Discover OpenAPI document"),
            DiagramNode("validate", "Validate Swagger 2 / OpenAPI 3"),
            DiagramNode("preview", "Preview operations", "ui"),
            DiagramNode("select", "Select operations", "ui"),
            DiagramNode("convert", "Convert selected services"),
            DiagramNode("commit", "Add to project and mark dirty", "ui"),
            DiagramNode("error", "Show actionable error", "sensitive"),
        ),
        (
            DiagramEdge("start", "fetch"), DiagramEdge("fetch", "discover"),
            DiagramEdge("discover", "validate"), DiagramEdge("validate", "preview", "valid"),
            DiagramEdge("validate", "error", "invalid"), DiagramEdge("preview", "select"),
            DiagramEdge("select", "convert", "one or more"), DiagramEdge("select", "preview", "none selected"),
            DiagramEdge("convert", "commit"),
        ),
        ("Nothing changes in the project until the user confirms a non-empty selection.", "Example credentials are never imported."),
    ),
    DiagramSpec(
        "state-machine-diagram",
        "Request and Flow Execution State Machine",
        "State Machine Diagram",
        "Defines visible execution states and terminal behavior used by the workbench, flow canvas, console, and cancellation controls.",
        (
            DiagramNode("idle", "idle"), DiagramNode("running", "running", "ui"),
            DiagramNode("success", "success"), DiagramNode("failed", "failed", "sensitive"),
            DiagramNode("cancelled", "cancelled", "sensitive"), DiagramNode("skipped", "skipped"),
            DiagramNode("blocked", "blocked", "sensitive"),
        ),
        (
            DiagramEdge("idle", "running", "Send / Run Flow"),
            DiagramEdge("running", "success", "2xx/accepted completion"),
            DiagramEdge("running", "failed", "validation or terminal error"),
            DiagramEdge("running", "cancelled", "AbortSignal"),
            DiagramEdge("failed", "running", "retryable and attempts remain"),
            DiagramEdge("failed", "blocked", "dependent flow node"),
            DiagramEdge("success", "idle", "reset"), DiagramEdge("cancelled", "idle", "reset"),
            DiagramEdge("blocked", "skipped", "cleanup policy"),
        ),
        ("Every terminal state is visible; there are no silent fallbacks.", "Flow nodes use the same named state vocabulary as persisted UI state."),
    ),
    DiagramSpec(
        "sequence-diagram",
        "Request Execution, Retry, Cancellation, and Diagnostics",
        "Sequence Diagram",
        "Shows chronological responsibility during request execution and the points where secrets are resolved, transported, and redacted.",
        (
            DiagramNode("user", "User", "external"), DiagramNode("app", "Workbench", "ui"),
            DiagramNode("runner", "ServiceRunner"), DiagramNode("redaction", "Redaction"),
            DiagramNode("native", "Tauri HTTP", "native"), DiagramNode("api", "REST Target", "external"),
        ),
        (
            DiagramEdge("user", "app", "Send"), DiagramEdge("app", "runner", "run(service, env, signal)"),
            DiagramEdge("runner", "redaction", "prepare safe diagnostics"), DiagramEdge("runner", "native", "attempt 1"),
            DiagramEdge("native", "api", "HTTP request"), DiagramEdge("api", "native", "retryable 5xx/transport error"),
            DiagramEdge("native", "runner", "typed failure"), DiagramEdge("runner", "native", "bounded retry"),
            DiagramEdge("api", "native", "response"), DiagramEdge("native", "runner", "typed response"),
            DiagramEdge("runner", "redaction", "sanitize result"), DiagramEdge("runner", "app", "visible result"),
        ),
        ("If the user cancels, the AbortSignal terminates the active attempt and suppresses further retries.", "Passwords, bearer tokens, API keys, cookies, and secret variables never enter visible diagnostics."),
    ),
    DiagramSpec(
        "communication-diagram",
        "Project Save, Backup, and Recovery Collaboration",
        "Communication Diagram",
        "Emphasizes numbered collaboration between project state, validation, redaction, concurrency protection, native atomic IO, and recovery backup.",
        (
            DiagramNode("ui", "Workbench", "ui"), DiagramNode("state", "Project State"),
            DiagramNode("schema", "Schema Validator"), DiagramNode("redact", "Project Redactor"),
            DiagramNode("guard", "Save Guard"), DiagramNode("native", "Native Persistence", "native"),
            DiagramNode("backup", "Recovery Backup", "external"), DiagramNode("file", "Project File", "external"),
        ),
        (
            DiagramEdge("ui", "state", "1 request snapshot"), DiagramEdge("state", "schema", "2 validate"),
            DiagramEdge("schema", "redact", "3 valid project"), DiagramEdge("redact", "guard", "4 safe payload"),
            DiagramEdge("guard", "native", "5 exclusive save"), DiagramEdge("native", "backup", "6 preserve previous"),
            DiagramEdge("native", "file", "7 atomic replace"), DiagramEdge("native", "ui", "8 success or explicit error"),
        ),
        ("Concurrent saves fail explicitly instead of racing.", "Recovery guidance identifies the backup and original paths without exposing project secrets."),
    ),
    DiagramSpec(
        "interaction-overview-diagram",
        "End-to-End API Debugging Interaction Overview",
        "Interaction Overview Diagram",
        "Combines the major interactions a developer uses to diagnose an unfamiliar API from import through a reproducible flow and safe evidence bundle.",
        (
            DiagramNode("open", "Open/Create Project", "ui"),
            DiagramNode("import", "Import Selected Operations", "ui"),
            DiagramNode("configure", "Configure Environment and Auth", "ui"),
            DiagramNode("request", "Run Single Request", "ui"),
            DiagramNode("inspect", "Inspect Response and Console", "ui"),
            DiagramNode("flow", "Compose and Run Flow", "ui"),
            DiagramNode("cancel", "Cancel or Retry", "ui"),
            DiagramNode("diag", "Export Redacted Diagnostics", "ui"),
            DiagramNode("save", "Save Project and Recovery Backup", "ui"),
        ),
        (
            DiagramEdge("open", "import"), DiagramEdge("import", "configure"),
            DiagramEdge("configure", "request"), DiagramEdge("request", "inspect"),
            DiagramEdge("inspect", "configure", "adjust variables"), DiagramEdge("inspect", "flow", "request works"),
            DiagramEdge("flow", "cancel", "slow/failing"), DiagramEdge("cancel", "inspect"),
            DiagramEdge("inspect", "diag", "need evidence"), DiagramEdge("diag", "save"),
        ),
        ("Each interaction node expands to a detailed activity or sequence view elsewhere in the atlas.", "The workflow remains useful without a hosted account or cloud workspace."),
    ),
    DiagramSpec(
        "timing-diagram",
        "Request Timeout, Retry, Backoff, and Cancellation Timing",
        "Timing Diagram",
        "Makes bounded execution behavior visible when diagnosing slow endpoints, retry policies, and user cancellation.",
        (
            DiagramNode("t0", "t0\nSend\nrunning"),
            DiagramNode("t1", "t1\nAttempt 1\ntransport failure"),
            DiagramNode("backoff", "t1..t2\nBackoff\ncancellable"),
            DiagramNode("t2", "t2\nAttempt 2\nrunning"),
            DiagramNode("cancel", "cancel event\nAbortSignal", "sensitive"),
            DiagramNode("terminal", "terminal\ncancelled or response"),
            DiagramNode("timeout", "timeout limit\nexplicit failure", "sensitive"),
        ),
        (
            DiagramEdge("t0", "t1", "attempt duration"), DiagramEdge("t1", "backoff", "retryable"),
            DiagramEdge("backoff", "t2", "delay elapsed"), DiagramEdge("t2", "terminal", "response"),
            DiagramEdge("backoff", "cancel", "user cancels"), DiagramEdge("t2", "cancel", "user cancels"),
            DiagramEdge("cancel", "terminal", "no more attempts"), DiagramEdge("t2", "timeout", "deadline reached"),
            DiagramEdge("timeout", "terminal", "typed timeout"),
        ),
        ("Backoff and timeout values come from the service and project settings.", "Cancellation wins over retry scheduling and produces a visible cancelled state."),
    ),
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr_text, fld_char2))


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.text = short_title
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    core = doc.core_properties
    core.title = short_title
    core.subject = "Relay Studio authoritative documentation"
    core.author = "Relay Studio Project"
    core.last_modified_by = "Relay Studio Project"
    core.keywords = "Relay Studio, documentation, developer tooling"


def add_cover(doc: Document, title: str, subtitle: str, audience: str, version_line: str = "Version 1.2 | Sprint 17 | July 2026") -> None:
    doc.add_paragraph("RELAY STUDIO", style="Heading 3").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle, style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph(f"Audience: {audience}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph(version_line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str], space_after: float = 4) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(space_after)
        p.add_run(item)


def add_steps(doc: Document, items: Iterable[str]) -> None:
    numbering = doc.part.numbering_part.element
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), "7")
    num.append(abstract_num_id)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        direct_num_id = OxmlElement("w:numId")
        direct_num_id.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, direct_num_id))
        p._p.get_or_add_pPr().append(num_pr)
        p.add_run(item)


def add_code(doc: Document, lines: Iterable[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shading)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_section(doc: Document, section: SectionSpec) -> None:
    heading = doc.add_heading(section.heading, level=1)
    heading.paragraph_format.page_break_before = section.page_break_before
    paragraphs = (section.paragraphs,) if isinstance(section.paragraphs, str) else section.paragraphs
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if section.bullets:
        add_bullets(doc, section.bullets)
    if section.steps:
        add_steps(doc, section.steps)
    if section.code:
        add_code(doc, section.code)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width_twips = [round(width * 1440) for width in widths]
    table_width = sum(width_twips)
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(table_width))
    tbl_indent = OxmlElement("w:tblInd")
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_indent.set(qn("w:w"), "120")
    tbl_pr.append(tbl_indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in width_twips:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_manual(filename: str, title: str, subtitle: str, audience: str, sections: Sequence[SectionSpec], callout: str, version_line: str = "Version 1.7 | Sprint 18C | July 2026", page_break_after_contents: bool = False, contents_space_after: float = 4) -> Path:
    doc = Document()
    configure_document(doc, title)
    add_cover(doc, title, subtitle, audience, version_line)
    add_callout(doc, "How to use this guide", callout)
    doc.add_heading("Contents", level=1)
    add_bullets(doc, (section.heading for section in sections), space_after=contents_space_after)
    if page_break_after_contents:
        doc.add_page_break()
    for section in sections:
        add_section(doc, section)
    path = WORD_DIR / filename
    doc.save(path)
    return path


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size, index=0)
    return ImageFont.load_default()


def node_positions(count: int, width: int, height: int) -> list[tuple[int, int]]:
    if count <= 6:
        cols = 3
    else:
        cols = 4
    rows = math.ceil(count / cols)
    margin_x, margin_y = 120, 190
    usable_w, usable_h = width - 2 * margin_x, height - margin_y - 120
    positions: list[tuple[int, int]] = []
    for index in range(count):
        row, col = divmod(index, cols)
        row_count = min(cols, count - row * cols)
        cell_w = usable_w / max(row_count, 1)
        x = int(margin_x + cell_w * (col + 0.5))
        y = int(margin_y + usable_h * (row + 0.5) / rows)
        positions.append((x, y))
    return positions


def render_preview(spec: DiagramSpec) -> Path:
    width, height = 1800, 1050
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 125), fill="#0B2545")
    draw.text((60, 28), spec.title, font=font(38, True), fill="white")
    draw.text((62, 78), spec.uml_type, font=font(21), fill="#DCE7F5")

    positions = node_positions(len(spec.nodes), width, height)
    centers = {node.key: positions[index] for index, node in enumerate(spec.nodes)}
    box_w, box_h = 310, 130
    fills = {
        "component": "#E8EEF5", "ui": "#DCEBFF", "native": "#E5E1F5",
        "external": "#F2F4F7", "sensitive": "#FCE8E6",
    }

    for edge in spec.edges:
        x1, y1 = centers[edge.source]
        x2, y2 = centers[edge.target]
        draw.line((x1, y1, x2, y2), fill="#5F6B7A", width=4)
        angle = math.atan2(y2 - y1, x2 - x1)
        tip = (x2 - math.cos(angle) * box_w * 0.48, y2 - math.sin(angle) * box_h * 0.48)
        left = (tip[0] - 18 * math.cos(angle - 0.55), tip[1] - 18 * math.sin(angle - 0.55))
        right = (tip[0] - 18 * math.cos(angle + 0.55), tip[1] - 18 * math.sin(angle + 0.55))
        draw.polygon((tip, left, right), fill="#5F6B7A")
        if edge.label:
            mx, my = (x1 + x2) // 2, (y1 + y2) // 2
            bbox = draw.textbbox((0, 0), edge.label, font=font(16))
            pad = 5
            draw.rectangle((mx - (bbox[2] - bbox[0]) // 2 - pad, my - 12 - pad, mx + (bbox[2] - bbox[0]) // 2 + pad, my + 12 + pad), fill="white")
            draw.text((mx, my), edge.label, anchor="mm", font=font(16), fill="#344054")

    for node, (x, y) in zip(spec.nodes, positions):
        x0, y0, x1, y1 = x - box_w // 2, y - box_h // 2, x + box_w // 2, y + box_h // 2
        draw.rounded_rectangle((x0, y0, x1, y1), radius=14, fill=fills[node.category], outline="#2E74B5", width=3)
        lines = node.label.split("\n")
        total_h = sum(draw.textbbox((0, 0), line, font=font(18 if i else 20, i == 0))[3] for i, line in enumerate(lines)) + 7 * (len(lines) - 1)
        current_y = y - total_h // 2
        for index, line in enumerate(lines):
            fnt = font(20 if index == 0 else 18, index == 0)
            bbox = draw.textbbox((0, 0), line, font=fnt)
            draw.text((x, current_y), line, anchor="ma", font=fnt, fill="#0B2545")
            current_y += bbox[3] - bbox[1] + 7

    draw.text((60, height - 55), "Relay Studio | Sprint 17 documentation", font=font(16), fill="#5F6B7A")
    path = PREVIEW_DIR / f"{spec.slug}.png"
    image.save(path, optimize=True)
    return path


def page_xml(spec: DiagramSpec) -> str:
    positions = node_positions(len(spec.nodes), 1800, 1050)
    centers = {node.key: positions[index] for index, node in enumerate(spec.nodes)}
    node_ids = {node.key: index + 1 for index, node in enumerate(spec.nodes)}
    shapes: list[str] = []
    for node, (x, y) in zip(spec.nodes, positions):
        pin_x = x / 180.0
        pin_y = (1050 - y) / 180.0
        shapes.append(
            f'<Shape ID="{node_ids[node.key]}" NameU="{escape(node.key)}" Type="Shape">'
            f'<Cell N="PinX" V="{pin_x:.3f}"/><Cell N="PinY" V="{pin_y:.3f}"/>'
            '<Cell N="Width" V="1.72"/><Cell N="Height" V="0.72"/>'
            '<Cell N="LineColor" V="#2E74B5"/><Cell N="FillForegnd" V="#E8EEF5"/>'
            '<Section N="Geometry" IX="0"><Row T="MoveTo" IX="1"><Cell N="X" V="0"/><Cell N="Y" V="0"/></Row>'
            '<Row T="LineTo" IX="2"><Cell N="X" V="1.72"/><Cell N="Y" V="0"/></Row>'
            '<Row T="LineTo" IX="3"><Cell N="X" V="1.72"/><Cell N="Y" V="0.72"/></Row>'
            '<Row T="LineTo" IX="4"><Cell N="X" V="0"/><Cell N="Y" V="0.72"/></Row>'
            '<Row T="LineTo" IX="5"><Cell N="X" V="0"/><Cell N="Y" V="0"/></Row></Section>'
            f'<Text>{escape(node.label)}</Text></Shape>'
        )
    connects: list[str] = []
    next_id = len(spec.nodes) + 1
    for edge in spec.edges:
        x1, y1 = centers[edge.source]
        x2, y2 = centers[edge.target]
        shapes.append(
            f'<Shape ID="{next_id}" NameU="connector-{next_id}" Type="Shape">'
            f'<Cell N="BeginX" V="{x1 / 180.0:.3f}"/><Cell N="BeginY" V="{(1050-y1) / 180.0:.3f}"/>'
            f'<Cell N="EndX" V="{x2 / 180.0:.3f}"/><Cell N="EndY" V="{(1050-y2) / 180.0:.3f}"/>'
            f'<Text>{escape(edge.label)}</Text></Shape>'
        )
        connects.append(f'<Connect FromSheet="{next_id}" FromCell="BeginX" ToSheet="{node_ids[edge.source]}" ToCell="PinX"/>')
        connects.append(f'<Connect FromSheet="{next_id}" FromCell="EndX" ToSheet="{node_ids[edge.target]}" ToCell="PinX"/>')
        next_id += 1
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<PageContents xmlns="http://schemas.microsoft.com/office/visio/2012/main">'
        f'<Shapes>{"".join(shapes)}</Shapes><Connects>{"".join(connects)}</Connects></PageContents>'
    )


def content_types(page_count: int) -> str:
    overrides = ''.join(
        f'<Override PartName="/visio/pages/page{i}.xml" ContentType="application/vnd.ms-visio.page+xml"/>'
        for i in range(1, page_count + 1)
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/visio/document.xml" ContentType="application/vnd.ms-visio.drawing.main+xml"/>'
        '<Override PartName="/visio/pages/pages.xml" ContentType="application/vnd.ms-visio.pages+xml"/>'
        '<Override PartName="/visio/masters/masters.xml" ContentType="application/vnd.ms-visio.masters+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        f'{overrides}</Types>'
    )


def build_vsdx(path: Path, specs: Sequence[DiagramSpec]) -> None:
    pages = ''.join(f'<Page ID="{i}" NameU="{escape(spec.uml_type)}" Name="{escape(spec.uml_type)}" r:id="rId{i}"/>' for i, spec in enumerate(specs, 1))
    page_rels = ''.join(
        f'<Relationship Id="rId{i}" Type="http://schemas.microsoft.com/visio/2010/relationships/page" Target="page{i}.xml"/>'
        for i in range(1, len(specs) + 1)
    )
    timestamp = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types(len(specs)))
        archive.writestr("_rels/.rels", '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="visio/document.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/></Relationships>')
        archive.writestr("docProps/core.xml", f'<?xml version="1.0" encoding="UTF-8"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>{escape(path.stem)}</dc:title><dc:creator>Relay Studio Project</dc:creator><dcterms:created xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:created><dcterms:modified xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:modified></cp:coreProperties>')
        archive.writestr("docProps/app.xml", '<?xml version="1.0" encoding="UTF-8"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Relay Studio Documentation Builder</Application><AppVersion>1.0</AppVersion></Properties>')
        archive.writestr("visio/document.xml", '<?xml version="1.0" encoding="UTF-8"?><VisioDocument xmlns="http://schemas.microsoft.com/office/visio/2012/main"><DocumentSettings/><Colors/></VisioDocument>')
        archive.writestr("visio/_rels/document.xml.rels", '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.microsoft.com/visio/2010/relationships/pages" Target="pages/pages.xml"/><Relationship Id="rId2" Type="http://schemas.microsoft.com/visio/2010/relationships/masters" Target="masters/masters.xml"/></Relationships>')
        archive.writestr("visio/pages/pages.xml", f'<?xml version="1.0" encoding="UTF-8"?><Pages xmlns="http://schemas.microsoft.com/office/visio/2012/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">{pages}</Pages>')
        archive.writestr("visio/pages/_rels/pages.xml.rels", f'<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{page_rels}</Relationships>')
        archive.writestr("visio/masters/masters.xml", '<?xml version="1.0" encoding="UTF-8"?><Masters xmlns="http://schemas.microsoft.com/office/visio/2012/main"/>')
        for index, spec in enumerate(specs, 1):
            archive.writestr(f"visio/pages/page{index}.xml", page_xml(spec))


def add_picture_with_alt(doc: Document, path: Path, alt_text: str, width: float = 6.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)


def build_uml_guide(previews: dict[str, Path]) -> Path:
    doc = Document()
    title = "Relay Studio UML Guide"
    configure_document(doc, title)
    add_cover(doc, title, "Fourteen complementary views of the application", "Developers, architects, QA, and product managers")
    add_callout(doc, "Deliverable", "Every diagram is available as an individual editable Visio file and as one page in Relay-Studio-UML-Atlas.vsdx. The preview in this guide is for reading and review.")
    doc.add_heading("How to Read the Atlas", level=1)
    doc.add_paragraph("Structural diagrams explain what exists. Behavioral diagrams explain how users and components collaborate over time. No single diagram is the architecture; use the smallest view that answers the current question.")
    add_table(doc, ("Category", "Diagram types", "Best question"), (
        ("Structural", "Class, Object, Component, Deployment, Package, Profile, Composite Structure", "What exists and how is it organized?"),
        ("Behavioral", "Use Case, Activity, State Machine, Sequence, Communication, Interaction Overview, Timing", "What happens, in what order, and under which conditions?"),
    ), (1.2, 3.3, 2.0))
    for index, spec in enumerate(DIAGRAMS, 1):
        doc.add_page_break()
        doc.add_heading(f"{index}. {spec.uml_type}: {spec.title}", level=1)
        doc.add_paragraph(spec.purpose)
        add_picture_with_alt(doc, previews[spec.slug], f"{spec.uml_type} preview: {spec.title}", width=6.2)
        p = doc.add_paragraph(f"Visio source: documentation/uml/visio/{spec.slug}.vsdx")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        add_bullets(doc, spec.notes)
    doc.add_heading("Maintenance Rules", level=1)
    add_bullets(doc, (
        "Keep shape labels aligned with current UI and code terminology.",
        "Update both the individual diagram and the corresponding page in the master atlas.",
        "Run the VSDX structural validator after any edit.",
        "Do not place credentials, real tokens, or private endpoint values in diagrams.",
        "Revise this guide when a diagram's purpose, assumptions, or authoritative source changes.",
    ))
    path = WORD_DIR / "Relay-Studio-UML-Guide.docx"
    doc.save(path)
    return path


def onboarding_sections() -> tuple[SectionSpec, ...]:
    return (
        SectionSpec("What You Are Joining", ("Relay Studio is a local-first Tauri desktop application for designing, importing, executing, and debugging REST requests and multi-step flows. It deliberately keeps project files, credentials, saved responses, and diagnostics under the developer's control.", "A successful maintainer protects three invariants: failures are explicit, secrets are redacted, and every changed behavior is verified both automatically and in the running application."), bullets=("Frontend: React 18, strict TypeScript, Vite, and React Flow.", "Native layer: Rust and Tauri 2 for filesystem, menus, dialogs, HTTP, and desktop lifecycle.", "Verification: Vitest, Testing Library, Playwright, Rust tests, clippy, llvm-cov, dependency audits, and secret scanning.")),
        SectionSpec("Workstation Setup", ("Use Node 22 and the stable Rust toolchain. Native builds also require the platform prerequisites documented by Tauri. Never place live credentials in the repository."), steps=("Clone the repository and verify the working tree is clean.", "Install JavaScript dependencies with npm ci.", "Install the stable Rust toolchain and platform-specific Tauri prerequisites.", "Run npm run check:types and cargo test --manifest-path src-tauri/Cargo.toml.", "Start browser development with npm run dev or native development with npm run tauri dev."), code=("npm ci", "npm run verify", "cargo test --manifest-path src-tauri/Cargo.toml", "npm run tauri dev")),
        SectionSpec("Repository Orientation", ("Begin at src/App.tsx for workbench orchestration, then move toward the smallest typed subsystem that owns the behavior. Avoid fixing domain behavior in presentation-only code."), bullets=("src/services — request execution, flows, OpenAPI import, saved responses, formatting, and diagnostics.", "src/project — model, schema validation, native/browser persistence, recovery, and project selection.", "src/shell — shared command model and native shell integration.", "src/lib — typed application errors and shared redaction.", "src-tauri/src — native command implementations, menus, filesystem, and HTTP boundaries.", "src/test and e2e — release-gate and user-workflow coverage.", "documentation — authoritative criteria, architecture, QA, and delivery records.")),
        SectionSpec("Your First Safe Change", ("Use test-driven development even for small application changes. Keep the diff focused and make the failure observable before implementation."), steps=("Write or update the narrowest test that expresses the requested behavior and an important failure path.", "Run the test and confirm it fails for the intended reason.", "Implement the smallest typed change; throw an explicit error rather than returning null or logging and continuing.", "Run the targeted test, then the representative suite.", "Start Relay Studio, take control of the changed workflow, inspect visible results and runtime logs, and repeat until they agree.", "Run coverage, lint, type checking, builds, security checks, and inspect git diff before handoff.")),
        SectionSpec("Debugging the React Workbench", ("Reproduce UI failures with a controlled project and note the active explorer item, tab, environment, dirty state, response dock, and console selection. State spread across these surfaces often explains an apparent rendering defect."), bullets=("Start with the nearest App or service test and reproduce the state transition directly.", "Use browser developer tools for component errors, network failures in browser mode, focus, and accessibility state.", "Inspect action handlers in App.tsx and follow typed data into the owning service module.", "For menu-only behavior, verify the shared shell command and native event path rather than duplicating logic.")),
        SectionSpec("Debugging Requests and Flows", ("Request execution is a typed pipeline: resolve variables, validate, build, transport, format, redact diagnostics, and update visible state. Flows add dependency ordering, mappings, branching, cleanup, cancellation, and node status."), bullets=("Reproduce with the smallest request and a controlled local endpoint.", "Check resolved URL, headers, query/path parameters, auth mode, timeout, retry policy, and active environment without printing secret values.", "For multipart files, confirm the user approved the exact service, field, local path, and destination origin in the current session; then confirm Tauri sees a readable regular file no larger than 25 MiB with a valid media type and no manual Content-Type header.", "If a reopened project has an empty file path, reselect or re-enter the fixture and approve it; saved projects deliberately do not retain local-file authority.", "Classify the failure as validation, authorization, transport, HTTP response, parsing, mapping, cancellation, or persistence.", "Inspect the visible console and exported diagnostics for a matching typed error and redacted values.", "For a flow, inspect the first non-success node and verify downstream blocked/skipped behavior."), page_break_before=True),
        SectionSpec("Debugging Native Tauri Behavior", ("A browser success does not prove native behavior. File dialogs, atomic writes, menu commands, packaged resources, proxy use, and native HTTP require a Tauri run."), bullets=("Read the Rust terminal output and browser/webview console together.", "Confirm the command is registered, included in the reviewed capability, and invoked with the expected payload.", "Test permission denial, missing files, malformed project data, and interrupted writes.", "Never broaden filesystem or network capabilities to hide a specific error."), page_break_before=True),
        SectionSpec("OpenAPI Import Failures", ("The importer accepts a Swagger UI page or direct Swagger/OpenAPI document, resolves bounded same-origin references, previews operations and review counts, and changes the project only after explicit selection."), bullets=("Separate URL discovery failures from definition parsing and external-reference failures.", "Verify Swagger discovery stops after the page fetch, displays the resolved destination, and retrieves it only after the explicit Load action; Cancel must send no second request.", "Verify every direct and referenced document against its actual final origin, not only its requested URL.", "Expect explicit failures for credential-bearing URLs, cross-origin redirects/references, circular, malformed, unreachable, excessively deep, or excessive-document reference graphs.", "Browser development mode deliberately blocks redirects before loading a document; enter the final definition URL explicitly or verify validated same-origin redirects in desktop mode.", "Check Swagger 2.0 versus OpenAPI 3.x server/base-path construction.", "Confirm PATCH, HEAD, OPTIONS, JSON, URL-encoded forms, and mixed text/file multipart forms map correctly.", "Confirm binary properties become empty file fields and examples do not import file paths, bearer tokens, passwords, API keys, cookies, or credentials.", "Test Select All, Clear, partial selection, Add Selected, Add and Save Selected, save cancellation, persistence/reopen, and duplicate ID generation."), page_break_before=True),
        SectionSpec("Persistence, Recovery, and Concurrent Saves", ("Project persistence deeply validates schema-v1 state, canonically redacts secret-bearing values, removes multipart file authority, protects concurrent saves, creates recovery backups, and uses native atomic replacement where available."), bullets=("Do not bypass nested schema validation to open malformed services, environments, flows, responses, imports, or settings.", "Older schema-v1 files may omit settings introduced later; verify those missing fields receive current typed defaults while explicitly malformed values still fail.", "Preserve the exact failing field path and provide recovery guidance.", "Inspect project and backup paths without printing content that may be sensitive.", "Confirm query/path/form canaries, URL userinfo, sensitive query values, and credential-shaped flow captures cannot survive save/reload.", "Simulate a second save and verify it fails explicitly instead of racing.", "Verify a restored backup reopens and remains valid."), page_break_before=True),
        SectionSpec("Diagnostics and Secret Safety", ("Treat logs, screenshots, saved responses, project exports, test fixtures, normalized errors, and diagnostics bundles as possible exfiltration paths."), bullets=("Use the shared canonical redaction utility; do not invent local masking logic.", "Test bearer tokens, passwords, API-key spelling variants, cookies, client secrets, URL userinfo, sensitive query values, custom auth headers, and secret variables.", "Reapply canonical redaction when reopening or comparing an artifact; never trust a persisted redacted flag by itself.", "Inspect the generated JSON bundle before sharing it.", "Never add a real credential to a unit test, snapshot, command, issue, or commit.")),
        SectionSpec("Test Selection and Release Gates", ("Start narrow, then expand in proportion to risk. Server/native Rust code must remain above 90 percent line coverage."), code=("npm test -- --run <test-file>", "npm run check:types", "npm run lint", "npm run test:coverage", "npm run test:component", "npm run build", "cargo test --manifest-path src-tauri/Cargo.toml", "npm run check:rust-coverage", "npm run test:e2e", "npm run check:dependencies", "npm run check:secrets"), bullets=("Live REST tests require a gitignored local file or protected CI secret.", "Ordinary PR/main validation reports a missing live target; beta packaging remains blocked until configured live acceptance passes.", "Package changes require platform installer and packaged-app verification."), page_break_before=True),
        SectionSpec("Common Failure Playbook", bullets=("npm ci dependency conflict — verify package.json and package-lock.json changed together and peer ranges support the selected version; do not use --force as a fix.", "TypeScript passes locally but fails in CI — confirm Node version, generated files, case-sensitive paths, and clean npm ci behavior.", "Rust command not found — confirm command registration, capability permission, and frontend/native name agreement.", "Request appears stuck — inspect timeout, active retry/backoff, AbortSignal ownership, proxy settings, and transport logs.", "Project will not open — preserve the file, inspect the typed schema error, and test recovery backup restoration.", "UI control has no effect — verify focus, disabled state, shared command routing, and stale selection state.")),
        SectionSpec("Junior Developer Exercises", steps=("Run the complete unit suite and locate the coverage report for one service module.", "Trace a Send Request click from App.tsx through ServiceRunner and the native transport boundary.", "Add a deliberately invalid local project fixture and explain the recovery guidance without weakening validation.", "Import two selected operations from a local OpenAPI fixture and confirm unselected operations are absent.", "Cancel a delayed request and verify the visible state, console, and diagnostics agree.", "Explain the class, component, sequence, and deployment diagrams to another developer."), page_break_before=True),
        SectionSpec("Definition of Done", bullets=("The requested behavior and failure modes have tests.", "Strict types cover parameters and return values.", "Errors are explicit and actionable.", "No secrets are committed or displayed.", "Representative tests, edge cases, coverage, lint, type checks, builds, and security checks pass.", "Relay Studio is running, the changed behavior is exercised interactively, and visible output plus runtime logs match expectations.", "Documentation and UML are updated when behavior or architecture changes.", "Commit messages include dev: jhaynes."), page_break_before=True),
    )


def architecture_sections() -> tuple[SectionSpec, ...]:
    return (
        SectionSpec("Executive Overview", ("Relay Studio is an installable local desktop workbench for developers who need to construct, import, execute, organize, and debug REST APIs without sending project data to a hosted workspace. Its differentiators are selective Swagger/OpenAPI import, visual flow composition, strong secret redaction, local project recovery, and desktop-native execution.", "The product is intentionally an unsigned personal developer tool. Paid Apple and Microsoft readiness programs, app stores, hosted collaboration, and enterprise deployment certification are outside the current goal."), page_break_before=True),
        SectionSpec("Product Capabilities", bullets=("Create, rename, delete, save, reopen, import, and recover local projects.", "Design REST services with GET, POST, PUT, PATCH, DELETE, HEAD, and OPTIONS plus parameters, headers, JSON/text/form bodies, multipart text/file parts, auth, timeout, retry, proxy, and environment variables.", "Import a Swagger UI page or direct OpenAPI/Swagger definition, explicitly review secondary destinations, resolve bounded same-origin external references, review counts, and select operations.", "Run and cancel individual requests, inspect formatted responses and final origin, and export redacted diagnostics with a request inventory.", "Compose visual flows with dependency edges, failure branches, protected destination variables, mappings, cleanup, and cancellation.", "Save redacted response artifacts, reopen them, and compare two artifacts for metadata plus structured JSON or raw-line differences.", "Package unsigned developer betas for macOS, Windows, and Linux."), page_break_before=True),
        SectionSpec("System Context", ("The desktop process hosts a React webview and a Rust Tauri layer. The user controls local project/backup/response files and configured REST targets. GitHub Actions consumes dependency metadata and protected live-test configuration only for quality, security, and packaging automation."), bullets=("Trust boundary 1: React/TypeScript to registered Tauri commands.", "Trust boundary 2: native process to filesystem and configured HTTP(S) endpoints.", "Trust boundary 3: CI to npm, crates.io, RustSec, GitHub artifacts, and protected secrets.", "Trust boundary 4: imported API definitions and project files as untrusted input."), page_break_before=True),
        SectionSpec("Frontend Architecture", ("App.tsx owns workbench composition and delegates domain behavior to typed modules. The shell includes native/shared commands, explorer, tabs, contextual tools, editors, inspector, response dock, and console. React Flow renders flow nodes and edges while service modules own graph validation and execution semantics."), bullets=("Keep presentation state close to the owning surface.", "Keep reusable domain behavior in src/services or src/project.", "Use shared typed errors and redaction utilities.", "Avoid global mutable state and UI-only business rules."), page_break_before=True),
        SectionSpec("Native Architecture", ("The Rust layer provides the desktop capabilities that cannot be trusted to browser-only adapters: atomic project IO, backups, response files, dialogs, menus, Help resources, close protection, and native HTTP execution. Structured multipart parts cross the existing request command boundary only after current-session approval; Rust reads the approved path at send time and streams the generated multipart form through Reqwest."), bullets=("Every command is explicitly registered.", "Native HTTP follows only same-origin redirects, stops after 10 requests, rejects cross-origin replay, and returns final URL identity to typed callers.", "Configured proxies apply validated domain, IP, CIDR, or wildcard bypass entries; malformed or port-specific entries fail explicitly.", "Multipart files must be regular files no larger than 25 MiB; malformed names, kinds, media types, missing paths, and manual multipart Content-Type headers fail explicitly.", "Saved `.json` and `.txt` responses are validated self-describing Relay envelopes; arbitrary raw text and embedded-path mismatches fail before content enters the response viewer.", "Capabilities remain least-privilege and are re-reviewed when commands or filesystem/network scope changes.", "Native errors cross the boundary as actionable failures.", "Packaged resources must be verified inside the installer, not only in development."), page_break_before=True),
        SectionSpec("Project Data Model", ("A Relay project has an explicit format identifier and schema version. It owns services, environments, flows, saved-response metadata, import sources, and settings. Services own request construction and auth; flows reference services by ID and add graph/mapping state. Multipart approval is ephemeral application state and never part of `.restproj`."), bullets=("Full nested schema validation occurs before project data enters the application.", "Missing settings in older schema-v1 files are migrated to typed current defaults; explicit invalid values identify their exact path and fail.", "Persistence and export canonically redact sensitive values and clear local file paths.", "Credential-shaped captured variables remain secret even when a mapping was mislabeled.", "Unknown or invalid schema shapes identify the exact field and produce recovery guidance.", "Concurrent saves and partial native writes are guarded."), page_break_before=True),
        SectionSpec("Request and Flow Execution", ("Single requests resolve environment variables, validate request construction, apply auth, use bounded retry/backoff, accept cancellation, format the response, and emit redacted console/diagnostic events. Flow execution extends that pipeline with graph ordering, success/failure edges, mappings, cleanup, node status, and a flow-level AbortSignal."), bullets=("Retry only classified retryable failures.", "Cancellation prevents future attempts and produces a terminal cancelled state.", "Blocked and skipped flow nodes remain visible.", "Mappings use explicit JSONPath and named variables; `baseUrl` is reserved and cannot be written from a response.", "Successful request metadata retains final URL identity for policy enforcement while the visible response summary exposes only final origin."), page_break_before=True),
        SectionSpec("OpenAPI Import", ("The importer treats supplied URLs, documents, and references as untrusted. It discovers definition URLs from Swagger UI pages without retrieving them, supports Swagger 2.0 and OpenAPI 3.x documents, resolves bounded same-origin external references, previews review metrics and operations, and converts only explicitly selected services."), bullets=("The project is unchanged until confirmation.", "Swagger UI page inspection presents the exact resolved secondary destination; only Load retrieves it and Cancel sends no request.", "Every root and external-document fetch revalidates actual final origin before parsing.", "Credential userinfo and literal sensitive query values are rejected before display or retrieval; safe `{{variable}}` placeholders remain supported.", "Add Selected keeps the project dirty; Add and Save Selected opens project save immediately.", "Canceling save keeps imported services dirty and editable.", "Imported service IDs are collision-safe.", "Relative servers, paths, and reference URLs resolve against the owning source document.", "Cross-origin, circular, malformed, unreachable, over-depth, and excessive-document graphs fail explicitly.", "Credentials and credential-like examples are replaced with safe variable placeholders.", "URL-encoded schemas import as text fields; multipart binary properties import as empty file fields with safe media types and never copy example local paths."), page_break_before=True),
        SectionSpec("Persistence and Recovery", ("Browser-mode persistence supports automated UI workflows. Native persistence adds platform file dialogs, atomic replacement, recovery backups, native recent-project paths, and packaged behavior. Saved responses are separate canonically redacted, self-describing artifacts referenced by validated metadata in the project; legacy raw `.txt` files require re-sending and saving a new artifact."), page_break_before=True),
        SectionSpec("Security Architecture", bullets=("Restrictive Tauri CSP and reviewed capability set.", "Canonical redaction for URLs, values, console, diagnostics, persistence, saved artifacts, comparison, normalized errors, export, and snapshots.", "Current-session multipart approval is bound to service, field, exact path, and destination origin; project save clears the path.", "Dependency advisories, license checks, secret scanning, clippy, coverage, and gitleaks in CI.", "No committed credentials or live target configuration.", "Only HTTP(S) live targets and non-placeholder credentials are accepted.", "Unsigned packages are clearly documented as developer betas."), page_break_before=True),
        SectionSpec("Build, Test, and Release Architecture", ("Release Gates run web quality, Rust quality, security, and optional live REST validation. Ordinary main validation can complete without protected live configuration, but beta packaging requires configured live REST acceptance before the macOS, Windows, and Linux matrix starts."), bullets=("Web gate: types, lint, coverage, component tests, production build, secret scan, Chromium/WebKit Playwright.", "Rust gate: tests, clippy with warnings denied, and at least 90 percent line coverage.", "Security gate: npm audit, license policy, cargo-deny, repository/artifact secret scanning, and gitleaks.", "Packaging gate: live REST preflight, platform build, verification, Rust tests, Tauri bundles, and artifact retention."), page_break_before=True),
        SectionSpec("Repository Structure", bullets=("src/App.tsx and src/styles.css — workbench UI and interaction orchestration.", "src/services — request/flow/import/response/diagnostic application services.", "src/project — persisted model, validation, project selection, backup, and IO adapters.", "src/shell — shared command and native shell contracts.", "src/lib — error and redaction foundations.", "src-tauri — Rust native application, Tauri configuration, capabilities, icons, and bundles.", "e2e and src/test — browser workflows, release contracts, and live REST acceptance.", ".github/workflows — release and packaging automation.", "documentation — authoritative directives, Word manuals, Visio UML, screenshots, and traceability."), page_break_before=True),
        SectionSpec("Architecture Decisions and Constraints", bullets=("Local-first files instead of a hosted workspace.", "Tauri instead of Electron for the desktop boundary.", "Strict typed application services instead of component-local request logic.", "Explicit errors instead of silent fallback.", "Selective import instead of automatic bulk project mutation.", "Shared redaction instead of per-surface masking.", "Unsigned personal-use packages instead of paid signing/readiness programs."), page_break_before=True),
        SectionSpec("Product Manager Handoff", ("A product manager should use the UML use-case and activity views to explain customer value, the component/deployment views to explain scope and risk, the Sprint Portfolio to understand delivered increments, and the Test and QA Manual to judge readiness."), bullets=("New scope must identify the user workflow, measurable success, failure modes, persistence impact, security impact, and platform impact.", "Any project schema change requires migration/recovery planning.", "Any new native command requires capability and platform review.", "Any new diagnostic field requires redaction analysis.", "Packaging changes require platform installers and packaged-app smoke tests."), page_break_before=True),
        SectionSpec("Documentation Map", bullets=("Developer Onboarding and Debugging Guide — build, trace, diagnose, test, and contribute.", "UML Guide and Visio Atlas — architecture and behavior views.", "Sprint Portfolio — delivery history, decisions, and roadmap.", "Test and QA Manual — current automated and human verification.", "Security, Platform, and Release Manual — trust boundaries, audits, packaging, and release policy.", "Retained Markdown directives — machine-readable build criteria, terminology, design constraints, and operative procedures."), page_break_before=True),
    )


def sprint_sections() -> tuple[SectionSpec, ...]:
    sections = [SectionSpec("Portfolio Purpose", ("This portfolio records the product's evolution without preserving every obsolete procedure. Detailed historical files remain available through Git history; current implementation and QA authority lives in the active manuals and retained directives.",))]
    sprint_summaries = (
        ("Sprint 0", "Approved the supplied reference screenshots and local-first REST/API product direction."),
        ("Sprint 1", "Defined the IDE-style shell, terminology, navigation, command placement, import workflow, and testing strategy."),
        ("Sprint 2", "Created the Tauri/React desktop foundation, visual shell, test harness, and initial acceptance gate."),
        ("Sprint 3", "Added project lifecycle, validation, dirty state, recent projects, native persistence, and recovery groundwork."),
        ("Sprint 4", "Implemented REST service design, typed request definitions, validation, and high coverage."),
        ("Sprint 5", "Added native request execution, response formatting, visible console diagnostics, and controlled failure testing."),
        ("Sprint 6", "Implemented redacted saved responses, native/browser file persistence, and reload workflows."),
        ("Sprint 7 / 7A", "Delivered visual flow composition, branching, ordering, validation, execution, and simplified workbench UX."),
        ("Sprint 8 / 8A / 8B", "Added mappings and variables, hardened flow authoring, and completed the compact desktop density pass."),
        ("Sprint 9A / 9B", "Unified shared/native shell commands, navigation, recent projects, inspector ownership, and status messaging."),
        ("Sprint 10A / 10B", "Completed interactive UI regression work plus macOS and Windows platform-specific audits and handoff."),
        ("Sprint 11", "Hardened configurable live REST acceptance with local-only credentials, explicit skips, target validation, and redaction assertions."),
        ("Sprint 12", "Added typed errors, retry/backoff, request/flow cancellation, project schema recovery, backups, save guards, and diagnostics export."),
        ("Sprint 13", "Enforced coverage, types, lint, component/UI/Rust tests, dependency/license audits, secret scanning, and release-gate policy."),
        ("Sprint 14", "Produced macOS, Windows, and Linux beta packages, added offline Help, implemented selective OpenAPI import, and improved explorer actions."),
        ("Sprint 15", "Creates the authoritative Word library, editable Visio UML atlas, onboarding/debugging guide, and curated documentation migration."),
        ("Sprint 16", "Implemented bounded external references, safe examples, additional HTTP methods and form bodies, saved-response comparison, import review, diagnostic inventory, and a Body-panel layout fix."),
        ("Sprint\u00a017", "Completed native multipart file workflows: safe OpenAPI binary-field import, typed text/file form rows, project persistence, browser boundary guidance, bounded local-file validation, and exact mixed-part transmission."),
        ("Sprint\u00a018A", "Completed the bounded review baseline: reconciled repeated discovery into 26 validated instances, recorded code and architecture deviations, assigned evidence-backed severity, and published the dependency-ordered remediation register for Sprints 18B-18E."),
        ("Sprint\u00a018B", "Closed six validated network/import findings with same-origin-only redirects, final response identity, proxy bypass enforcement, explicit Swagger secondary-destination review, final-origin revalidation, credential-free import URLs, and protected flow destinations."),
        ("Sprint\u00a018C", "Closed eight local-file, persistence, project-integrity, flow-secret, and redaction findings with ephemeral origin-bound multipart approval, cleared persisted paths, self-describing response envelopes, deep schema-v1 validation, and canonical redaction across public output boundaries."),
        ("Sprint\u00a018D (planned)", "Protect flow destination and secret integrity, correct branch semantics, and bound OpenAPI and saved-response comparison resource use."),
        ("Sprint\u00a018E (planned)", "Reduce CI secret exposure, close installer, lockfile, and OOXML scanning gaps, run all applicable readiness gates, and publish the final redacted decision."),
    )
    for name, summary in sprint_summaries:
        if "(planned)" in name:
            bullets = ("Status: planned; no implementation evidence is claimed.", "Detailed objectives and acceptance criteria remain in the retained sprint plans.")
        elif "(in progress)" in name:
            bullets = ("Status: in progress; review evidence is not yet a final readiness decision.", "Detailed objectives and acceptance criteria remain in the retained sprint plans.")
        else:
            bullets = ()
        sections.append(
            SectionSpec(
                name,
                (summary,),
                bullets=bullets,
                page_break_before=name in {"Sprint 7 / 7A", "Sprint\u00a018C"},
            )
        )
    sections.append(SectionSpec("Standing Product Decisions", bullets=("Local-first developer tool; no hosted account required.", "Explicit operation selection during API import.", "Strict typing, explicit errors, test-driven changes, and at least 90 percent native/server coverage.", "No committed secrets and shared redaction across every output surface.", "Unsigned personal-use packages; no paid Apple or Microsoft readiness programs.", "Every application change requires automated and interactive verification.")))
    return tuple(sections)


def test_sections() -> tuple[SectionSpec, ...]:
    return (
        SectionSpec("QA Strategy", ("Test current product behavior, not sprint numbers. Begin with deterministic unit and component tests, then exercise complete browser/native workflows, and reserve live REST and packaged-platform checks for configured acceptance environments."), bullets=("Expected path and important negative paths.", "Typed error content and visible recovery guidance.", "Cancellation, retry, timeout, and interrupted persistence.", "Secret redaction in every output surface.", "Platform and packaged-resource behavior.")),
        SectionSpec("Preflight", steps=("Confirm a clean working tree and record the commit under test.", "Use Node 22 and supported Rust toolchains.", "Run npm ci from the repository root.", "Keep live REST configuration in a gitignored file or protected secret.", "Use synthetic endpoints and credentials only.", "Record OS, package type, browser engine, and test evidence location.")),
        SectionSpec("Automated Quality Suite", code=("npm run check:types", "npm run lint", "npm run test:coverage", "npm run test:component", "npm run build", "cargo test --manifest-path src-tauri/Cargo.toml", "npm run check:rust", "npm run check:rust-coverage", "npm run test:e2e"), bullets=("All commands must exit successfully.", "Rust/native coverage must remain at least 90 percent.", "Coverage exclusions require explicit review and must not hide application logic."), page_break_before=True),
        SectionSpec("Security and Dependency Gates", code=("npm audit --audit-level=high", "npm run check:licenses", "cargo deny --manifest-path src-tauri/Cargo.toml check", "npm run check:secrets"), bullets=("No high-risk dependency finding may be ignored silently.", "No unapproved dependency license.", "No committed credential or secret-bearing generated artifact.", "Tauri CSP, capabilities, and commands remain reviewed.")),
        SectionSpec("Desktop Launch and Shell", steps=("Start Relay Studio in native development mode.", "Verify the activity rail, explorer, tabs, request composer, inspector, response dock, console, status, and native menus.", "Open the command palette and execute representative project/view commands.", "Resize through small, medium, and large desktop widths and verify no inaccessible controls or overlap.", "Inspect the webview console and Rust terminal for unhandled errors."), page_break_before=True),
        SectionSpec("Project Lifecycle", bullets=("Create a project, edit it, observe dirty state, save, close, and reopen.", "Exercise Save As, Recent Projects, missing recent files, and unsaved-close prompts.", "Open an older schema-v1 project with partial settings and verify current defaults are applied; then provide an explicitly malformed value and verify actionable recovery guidance.", "Simulate interrupted/native write behavior and restore the recovery backup.", "Trigger concurrent saves and verify the second operation fails explicitly."), page_break_before=True),
        SectionSpec("Service Design and Execution", bullets=("Edit method, URL, headers, query/path parameters, body, auth, timeout, retry, proxy, and environment values.", "For multipart, create text and file fields, save and reopen the project, and verify field kind plus media type persist while the local path is empty and approval is absent.", "Enter a synthetic path and confirm Send is blocked until approval; approve it, change origins, and confirm approval is invalidated before native file access.", "Send a controlled native multipart request after approval and verify field text, filename, media type, and exact file bytes at the receiver.", "Verify browser mode rejects local-file sends before Fetch and native mode rejects empty, missing, directory, oversized, malformed, or conflicting multipart inputs actionably.", "Verify a same-origin redirect succeeds and reports final origin in desktop mode; a host/port change, missing Location, or redirect loop fails before an unreviewed receiver is contacted. Browser development mode blocks all redirects and tells the developer to enter the final URL explicitly.", "Configure a direct receiver plus proxy and verify valid bypass entries route directly while malformed/port-specific entries fail explicitly.", "Run a successful request and inspect status, timing, content type, size, formatted body, final origin, console, and diagnostics.", "Exercise validation error, DNS/connection failure, timeout, malformed JSON, retryable failure, terminal HTTP response, and response-format fallback.", "Cancel a controlled delayed request and verify no retry continues.", "Rename and delete explorer items through their context menu."), page_break_before=True),
        SectionSpec("Selective OpenAPI Import", steps=("Choose File > Import API Definition.", "Enter a direct OpenAPI/Swagger URL and a Swagger UI page URL in separate cases.", "For Swagger UI, verify Inspect retrieves only the page and displays the resolved definition destination. Cancel and prove no secondary request; inspect again and choose Load.", "Verify definition discovery, preview metadata, operation count, external-document count, form-body count, and deprecated count.", "Use Select All and Clear, then choose a strict subset of operations.", "Use Add Selected to review before saving, then repeat with Add and Save Selected and confirm the project save dialog opens.", "Save and reopen the `.restproj` file; confirm only selected services persist. Cancel a second save and confirm imported services remain dirty and editable.", "Inspect generated PATCH/HEAD/OPTIONS methods, path and parameters, safe JSON examples, URL-encoded fields, and multipart text/file fields; binary properties must have empty paths and safe media types.", "Assign a synthetic local file and save/reopen; confirm its path and authority were removed, then re-enter it, approve the destination in desktop mode, and inspect the receiver. Repeat in browser mode and confirm the desktop-only error.", "Repeat credential-bearing URL, final-origin change, invalid document, zero-selection, duplicate-operation, credential-like example, example local-path, cross-origin, circular, malformed, unreachable, over-depth, and excessive-document cases."), page_break_before=True),
        SectionSpec("Saved Response Comparison", steps=("Save two redacted responses for the same or related developer workflow.", "Select exactly two artifacts in Saved Responses; confirm a third selection is disabled.", "Compare and inspect status, status text, duration, content type, and size changes.", "For JSON, verify added, removed, changed, and unchanged path counts; for non-JSON or malformed JSON, verify raw line changes.", "Delete or move one artifact and confirm the missing-file error identifies the exact local path without exposing response secrets."), page_break_before=True),
        SectionSpec("Flow Authoring and Execution", bullets=("Create a flow from services and verify node positions, ordering, success/failure edges, mappings, and cleanup affordances.", "Attempt a response mapping named `baseUrl` with case/whitespace variations and confirm validation blocks execution with environment-edit guidance.", "Run success, failure branch, blocked dependency, mapping failure, retry, and cleanup scenarios.", "Cancel a delayed flow and confirm the active node and flow become cancelled while future nodes do not run.", "Inspect grouped console output and flow summary for consistent status."), page_break_before=True),
        SectionSpec("Saved Responses and Diagnostics", bullets=("Save JSON and raw responses as self-describing Relay envelopes, verify canonical redaction, reopen them, and handle missing/permission-denied paths.", "Rename arbitrary local text to `.txt` and confirm it is rejected with guidance; substitute an artifact at a mismatched path and confirm it is rejected before display.", "Import an artifact that claims redacted state but contains synthetic URL/body canaries and confirm save, reopen, and comparison reapply redaction.", "Export a diagnostics bundle after a request and flow run.", "Inspect JSON directly for bearer tokens, passwords, API-key variants, cookies, client secrets, URL userinfo/query values, custom auth, and secret variables.", "Verify typed error codes, event ordering, request IDs, and useful non-secret context remain."), page_break_before=True),
        SectionSpec("Live REST Acceptance", ("Live acceptance uses configurable admin, standard, and restricted credentials against a user-selected HTTP(S) target. Missing local configuration skips explicitly; future beta packaging is blocked without protected configuration."), code=("RELAY_LIVE_REST_CONFIG=/absolute/gitignored/live-rest.json npm run test:live-rest",), bullets=("Health and authentication behavior.", "Authenticated reads and chained flows.", "Role denial and restricted-user boundaries.", "Timeout and malformed response behavior.", "Redaction of passwords, bearer tokens, API keys, and cookies."), page_break_before=True),
        SectionSpec("Packaged Platform Regression", bullets=("Install and launch the native artifact.", "Verify menus, title bar, caption controls, keyboard commands, file dialogs, Help, OpenAPI import, save/reopen, request/flow execution, diagnostics, and close protection.", "Verify platform appearance and accessibility expectations.", "Confirm offline Help is bundled.", "Windows: continue using the retained packaged-Windows QA script for installer, breakpoint, high-contrast, multi-monitor, persistence, and keyboard checks.", "macOS/Linux: record package type, permission behavior, filesystem paths, webview runtime, and known unsigned-package warnings."), page_break_before=True),
        SectionSpec("Evidence and Defect Record", bullets=("Commit and package version.", "OS, architecture, runtime, browser engine, and configuration class.", "Exact steps, expected result, actual result, screenshot/log locations, and secret-safety confirmation.", "Severity, reproducibility, affected subsystem, and smallest failing test.", "Retest result and final runtime-log inspection."), page_break_before=True),
        SectionSpec("Release Decision", bullets=("All required automated gates pass.", "Configured live REST passes for beta packaging; a skip is not release evidence.", "Changed behavior is exercised in the running native app.", "Required platform packages are produced and smoke tested.", "No P0/P1 defects, secret leakage, unhandled exceptions, or unexplained coverage regressions.", "Known unsigned-package limitations are documented."), page_break_before=True),
    )


def security_sections() -> tuple[SectionSpec, ...]:
    return (
        SectionSpec("Security Posture", ("Relay Studio reduces hosted-data exposure by remaining local-first, but it processes credentials, arbitrary URLs, imported definitions, project files, and response data. The primary risks are secret leakage, over-broad native capability, unsafe persistence, malicious/unexpected input, and misleading release evidence.")),
        SectionSpec("Secret Classification and Redaction", bullets=("Bearer and OAuth tokens, passwords, API keys across case/underscore/hyphen variants, cookies, client secrets, custom auth headers, URL userinfo, sensitive query values, credential-shaped response captures, and secret-marked variables are sensitive.", "Apply canonical redaction to console events, diagnostics, project export/persistence, saved responses, comparison, screenshots/fixtures, normalized errors, and snapshots before output is marked redacted.", "Preserve useful names, origins, non-sensitive query values, and structure while replacing the credential value consistently.", "Never trust a persisted redacted flag without reapplying canonical redaction.", "Never log a secret before redaction and never use real credentials in tests.")),
        SectionSpec("Tauri Boundary", bullets=("Restrictive CSP with object execution disabled and explicit script/connect sources.", "Reviewed capability file and registered commands only.", "Atomic writes, backups, and full nested schema validation for project persistence.", "Project save clears multipart paths; current-session approval binds service, field, path, and destination origin before native file access.", "Saved response commands require validated self-describing envelopes and matching embedded/requested paths; arbitrary raw `.txt` files never receive trusted caller metadata.", "Native HTTP follows only same-origin redirects, rejects cross-origin replay before headers can move, returns final identity, and stops redirect loops after 10 requests.", "Browser development mode uses manual Fetch redirects and rejects every redirect before replay; developers enter the final URL explicitly or use desktop mode for validated same-origin redirects.", "Configured proxies apply a validated bypass list instead of silently ignoring it.", "Re-review when adding commands, plugins, filesystem scope, remote content, updater behavior, or shell execution.")),
        SectionSpec("Dependency and Supply Chain", bullets=("npm audit blocks high-severity advisories.", "cargo-deny checks RustSec advisories, bans, sources, and licenses.", "The npm license script enforces the approved policy.", "Gitleaks and the project scanner cover repository and generated bundle paths.", "Dependency metadata may be sent to npmjs.org, crates.io, and RustSec; source code and secrets may not.")),
        SectionSpec("CI and Release Policy", ("PR and main Release Gates run deterministic quality and security validation. When live REST configuration is missing they report the omission without making ordinary CI permanently red. Beta-tag and manual Package Beta runs require the protected live configuration and complete acceptance before the platform matrix starts."), bullets=("Protected configuration is base64-decoded only into the runner temporary directory.", "Do not print, artifact, or cache the decoded file.", "Artifacts are unsigned developer betas retained for a bounded period.", "Signing/notarization deferral is intentional and visible in release notes.")),
        SectionSpec("Platform Packaging", bullets=("macOS: DMG, native menus, file dialogs, close protection, offline Help, unsigned warning expectations.", "Windows: NSIS and MSI, title/caption behavior, WebView2, high contrast, multi-monitor, native paths, and unsigned warning expectations.", "Linux: DEB and AppImage, WebKitGTK dependencies, native paths, and permission handling.", "Every packaging change requires installer reconstruction and packaged-app smoke testing.")),
        SectionSpec("Platform Audit Summary", ("The historical Apple HIG and Windows desktop audits established the native shell, dense workbench, menu parity, close protection, responsive breakpoints, accessibility appearance, and evidence requirements. Their enduring requirements are carried into the retained platform backlog and Test and QA Manual; dated findings remain available through Git history.")),
        SectionSpec("Incident and Diagnostics Handling", steps=("Preserve the original project and response artifacts without sharing them broadly.", "Reproduce with synthetic values if possible.", "Export a diagnostics bundle and inspect it locally for secret leakage.", "Record the typed error, timestamps, operation IDs, environment class, and platform without copying credentials.", "Fix through tests, run security/coverage gates, and verify in the native app.")),
        SectionSpec("Known Risks and Deferrals", bullets=("Unsigned packages can trigger platform trust warnings.", "Cross-origin HTTP redirects are intentionally rejected; the developer must enter the final destination explicitly.", "External OpenAPI references are intentionally same-origin and bounded; cross-origin graphs require the API owner to consolidate or proxy definitions locally.", "Multipart files are authorized only for the current session and destination origin, read at native send time, and limited to 25 MiB. Raw application/octet-stream bodies, streaming uploads, and browser file sends remain deferred.", "Legacy raw `.txt` response files require the developer to re-send and save a new self-describing artifact.", "User-configured endpoints may return hostile or very large data; response limits and formatting remain assigned to Sprint 18D.", "Local project files are only as protected as the user's filesystem account.", "There is no hosted collaboration, public updater, store distribution, or enterprise deployment certification.")),
        SectionSpec("Release Checklist", bullets=("Exact candidate commit identified.", "Web, Rust, security, and configured live REST gates pass.", "No high-risk dependency or secret-scanning finding.", "Required platform installers built and smoke tested.", "Help and documentation correspond to the packaged commit.", "Release notes identify unsigned status and known limitations.", "Diagnostics and artifacts inspected for sensitive content.")),
    )


def create_traceability() -> dict[str, object]:
    markdown_files = sorted(
        str(path.relative_to(ROOT))
        for base in (DOC_ROOT, ROOT / "audits")
        if base.exists()
        for path in base.rglob("*.md")
    )
    if (ROOT / "design-qa.md").exists():
        markdown_files.append("design-qa.md")
    entries = []
    for source in sorted(set(markdown_files)):
        if source in RETAINED_MARKDOWN:
            destination = "Retained authoritative Markdown; referenced by the Word library"
            action = "retain"
        elif "test" in source.lower() or "qa" in source.lower() or "acceptance" in source.lower():
            destination = "documentation/word/Relay-Studio-Test-and-QA-Manual.docx"
            action = "consolidate-remove"
        elif "audit" in source.lower() or "platform" in source.lower() or "release" in source.lower():
            destination = "documentation/word/Relay-Studio-Security-Platform-and-Release-Manual.docx"
            action = "consolidate-remove"
        elif "sprint" in source.lower() or "implementation-status" in source.lower():
            destination = "documentation/word/Relay-Studio-Sprint-Portfolio.docx"
            action = "consolidate-remove"
        else:
            destination = "documentation/word/Relay-Studio-Technical-Architecture-and-Product-Handoff.docx"
            action = "consolidate-remove"
        entries.append({"source": source, "action": action, "destination": destination})
    return {
        "schemaVersion": 1,
        "generatedAt": datetime.now(timezone.utc).replace(microsecond=0).isoformat(),
        "authority": {
            "word": "Narrative onboarding, architecture, UML guidance, sprint summaries, QA consolidation, and release operations",
            "markdown": "Build criteria, engineering policy, design constraints, and operative test procedures listed as retained",
        },
        "entries": entries,
    }


def main() -> None:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    VISIO_DIR.mkdir(parents=True, exist_ok=True)
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)

    portfolio_only = "--portfolio-only" in sys.argv
    manuals_only = "--manuals-only" in sys.argv or portfolio_only
    previews = {spec.slug: PREVIEW_DIR / f"{spec.slug}.png" for spec in DIAGRAMS}
    if not manuals_only:
        previews = {spec.slug: render_preview(spec) for spec in DIAGRAMS}
        for spec in DIAGRAMS:
            build_vsdx(VISIO_DIR / f"{spec.slug}.vsdx", (spec,))
        build_vsdx(VISIO_DIR / "Relay-Studio-UML-Atlas.vsdx", DIAGRAMS)

    if portfolio_only:
        build_manual(
            "Relay-Studio-Sprint-Portfolio.docx",
            "Relay Studio Sprint Portfolio",
            "Delivery history, decisions, evidence, and forward plan",
            "Product managers, maintainers, and technical leadership",
            sprint_sections(),
            "This is a curated history. Git retains deleted detail; active criteria live in the retained Markdown directives and current manuals.",
            "Version 1.7 | Sprint 18C | July 2026",
            True,
            2,
        )
        return

    build_manual(
        "Relay-Studio-Developer-Onboarding-and-Debugging-Guide.docx",
        "Relay Studio Developer Onboarding and Debugging Guide",
        "A practical path from first checkout to confident diagnosis",
        "Newly hired junior developers and maintainers",
        onboarding_sections(),
        "Read sequentially during onboarding, then return to the subsystem playbooks while debugging. Commands and paths are verified against the repository through Sprint 18C.",
    )
    build_manual(
        "Relay-Studio-Technical-Architecture-and-Product-Handoff.docx",
        "Relay Studio Technical Architecture and Product Handoff",
        "Product intent, system design, ownership, and decision context",
        "Product managers, technical leaders, developers, and QA",
        architecture_sections(),
        "Use the executive and product sections for stakeholder briefings; use the architecture and repository sections for scope, risk, and ownership decisions.",
    )
    if not manuals_only:
        build_uml_guide(previews)
    build_manual(
        "Relay-Studio-Sprint-Portfolio.docx",
        "Relay Studio Sprint Portfolio",
        "Delivery history, decisions, evidence, and forward plan",
        "Product managers, maintainers, and technical leadership",
        sprint_sections(),
        "This is a curated history. Git retains deleted detail; active criteria live in the retained Markdown directives and current manuals.",
        "Version 1.7 | Sprint 18C | July 2026",
        True,
        2,
    )
    build_manual(
        "Relay-Studio-Test-and-QA-Manual.docx",
        "Relay Studio Test and QA Manual",
        "Current behavior-based verification for browser, native, live, and packaged workflows",
        "Developers, QA engineers, and release maintainers",
        test_sections(),
        "Select the smallest applicable section for a change, then run the full release checks required by risk. Obsolete sprint scripts are intentionally not reproduced.",
    )
    build_manual(
        "Relay-Studio-Security-Platform-and-Release-Manual.docx",
        "Relay Studio Security, Platform, and Release Manual",
        "Trust boundaries, redaction, audits, packaging, and release decisions",
        "Maintainers, security reviewers, QA, and product leadership",
        security_sections(),
        "Treat this manual and the retained Markdown security directives together: the Markdown policies control implementation; this manual explains their system and release context.",
    )

    if not TRACEABILITY_PATH.exists():
        TRACEABILITY_PATH.write_text(json.dumps(create_traceability(), indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
